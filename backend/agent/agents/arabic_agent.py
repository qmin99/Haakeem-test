import logging
import base64
import json
import re
from livekit.agents import Agent
from livekit.plugins import groq
from livekit.plugins.azure import TTS as AzureTTS
from livekit.plugins.azure import STT as AzureSTT


class ArabicAgent(Agent):
    def __init__(self) -> None:
        super().__init__(
            instructions=(
                "اسمُك حَكيم، مساعد قانوني ذكي من Binfin8. "
                "تتكلم دائمًا بلهجة المتحدث إن تعرّفت عليها من كلامه. إن لم تتعرّف على لهجته فاختر لهجة خليجية طبيعية وواضحة (سعودية/قطرية) بدل العربية الفصحى. "
                "كن عمليًا ومباشرًا، وجّه المستخدم بخطوات واضحة ومبسطة، واطلب أي معلومات ناقصة بدقة. "
                "عند الرد على الملفات أو الأسئلة، قدّم ملخصًا موجزًا ثم نقاطًا قانونية دقيقة وقابلة للتنفيذ. "
                "تجنّب الجُمل الطويلة؛ اجعل الجمل قصيرة وسهلة الفهم، وابتعد عن المصطلحات المعقدة إن وُجد بديل شعبي واضح."
            ),
            stt=AzureSTT(language="ar-SA"),
            llm=groq.LLM(model="allam-2-7b"),
            tts=AzureTTS(voice="ar-OM-AbdullahNeural", language="ar-OM"),
        )

    async def on_enter(self):
        await self.session.generate_reply(
            instructions="السلام عليكم! أنا حَكيم، مساعدك القانوني. وش تبيني أساعدك فيه اليوم؟",
            allow_interruptions=True,
        )

    # async def on_final_transcription(self, text: str) -> str:
    #     """Filter transcription to keep only Arabic text and remove English words"""
    #     logger = logging.getLogger("multi-agent-ptt")
    #     logger.info(f"🔍 [ARABIC FILTER] Input: '{text}'")
        
    #     filtered_text = self._filter_arabic_only(text)
        
    #     logger.info(f"🔍 [ARABIC FILTER] Output: '{filtered_text}'")
    #     logger.info(f"🔍 [ARABIC FILTER] Filtered out: {len(text.split()) - len(filtered_text.split())} words")
        
    #     return filtered_text

    # async def on_user_speech_committed(self, user_input):
    #     """Override user speech committed to filter Arabic only"""
    #     logger = logging.getLogger("multi-agent-ptt")
        
    #     # Get the transcription text
    #     if hasattr(user_input, 'transcript') and user_input.transcript:
    #         original_text = user_input.transcript
    #         logger.info(f"🔍 [ARABIC SPEECH FILTER] Input: '{original_text}'")
            
    #         filtered_text = self._filter_arabic_only(original_text)
    #         logger.info(f"🔍 [ARABIC SPEECH FILTER] Output: '{filtered_text}'")
            
    #         # If no Arabic content, skip processing
    #         if not filtered_text.strip():
    #             logger.info("🔍 [ARABIC SPEECH FILTER] No Arabic content - skipping")
    #             return
            
    #         # Modify the transcript
    #         user_input.transcript = filtered_text
        
    #     # Call parent implementation
    #     await super().on_user_speech_committed(user_input)

    # def _filter_arabic_only(self, text: str) -> str:
    #     """Remove English words and keep only Arabic text"""
    #     import re
        
    #     # Split text into words
    #     words = text.split()
    #     filtered_words = []
        
    #     for word in words:
    #         # Remove punctuation for analysis
    #         clean_word = re.sub(r'[^\w\s]', '', word)
            
    #         # Skip empty words
    #         if not clean_word.strip():
    #             continue
                
    #         # Check if word contains Arabic characters
    #         if self._contains_arabic(clean_word):
    #             filtered_words.append(word)
    #         else:
    #             # Skip English words - don't add them
    #             continue
        
    #     # Return filtered text
    #     result = ' '.join(filtered_words)
        
    #     # If no Arabic words found, return empty string to avoid processing
    #     return result.strip()
    
    # def _contains_arabic(self, text: str) -> bool:
    #     """Check if text contains Arabic characters"""
    #     arabic_pattern = r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]'
    #     return bool(re.search(arabic_pattern, text))

    async def _file_received(self, reader, participant_identity):
        logger = logging.getLogger("multi-agent-ptt")
        stream_info = reader.info
        logger.info("📄 [Arabic] received byte stream: %s (%s)", stream_info.name, stream_info.mime_type)
        file_bytes = bytearray()
        async for chunk in reader:
            file_bytes.extend(chunk)
        await self._file_received_fallback(bytes(file_bytes), stream_info, participant_identity)

    async def _process_file(self, file_bytes, stream_info):
        logger = logging.getLogger("multi-agent-ptt")
        mime_type = stream_info.mime_type
        file_name = stream_info.name

        try:
            if mime_type == "text/plain":
                try:
                    content = file_bytes.decode("utf-8")
                    return content
                except UnicodeDecodeError:
                    for enc in ["latin-1", "cp1252", "iso-8859-1"]:
                        try:
                            content = file_bytes.decode(enc)
                            return content
                        except UnicodeDecodeError:
                            continue
                    return None
            elif mime_type == "application/pdf":
                try:
                    import io
                    import PyPDF2
                    pdf_stream = io.BytesIO(file_bytes)
                    pdf_reader = PyPDF2.PdfReader(pdf_stream)
                    extracted_text = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    if extracted_text.strip():
                        full_content = f"مستند PDF: {file_name}\nالمحتوى:\n{extracted_text}"
                        return full_content
                    return (
                        f"استلمت مستند PDF '{file_name}' لكن ما قدرت أستخرج نص واضح (يمكن يكون صور أو مشفّر)."
                    )
                except Exception as e:
                    return f"في مشكلة أثناء معالجة ملف PDF '{file_name}': {str(e)}"
            elif mime_type.startswith("image/"):
                image_b64 = base64.b64encode(file_bytes).decode("utf-8")
                return (
                    f"صورة '{file_name}' تم استلامها ({mime_type}, {len(file_bytes)} بايت). "
                    f"بيانات Base64 جاهزة للتحليل البصري: {image_b64[:100]}...[مقتطف]"
                )
            elif mime_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ]:
                try:
                    import io
                    from docx import Document
                    doc_stream = io.BytesIO(file_bytes)
                    doc = Document(doc_stream)
                    extracted_text = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            extracted_text += paragraph.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    extracted_text += cell.text + " "
                            extracted_text += "\n"
                    if extracted_text.strip():
                        full_content = f"مستند Word: {file_name}\nالمحتوى:\n{extracted_text}"
                        return full_content
                    return f"استلمت مستند Word '{file_name}' لكنه بدون نص واضح."
                except Exception as e:
                    return f"صار خطأ أثناء معالجة مستند Word '{file_name}': {str(e)}"
            elif mime_type == "application/json":
                try:
                    content = json.loads(file_bytes.decode("utf-8"))
                    json_content = f"ملف JSON '{file_name}'\nالمحتوى:\n{json.dumps(content, indent=2, ensure_ascii=False)}"
                    return json_content
                except Exception:
                    return None
            else:
                return (
                    f"استلمت ملف '{file_name}' بنوع غير مدعوم '{mime_type}'. الأنواع المدعومة: نصوص، PDF، صور، وملفات Word."
                )
        except Exception as e:
            logger.error("❌ [Arabic] processing error: %s", e, exc_info=True)
            return None

    async def _file_received_fallback(self, file_bytes, stream_info, participant_identity):
        logger = logging.getLogger("multi-agent-ptt")
        logger.info("📄 [Arabic] fallback upload: %s (%s)", stream_info.name, stream_info.mime_type)
        try:
            file_content = await self._process_file(file_bytes, stream_info)
            if file_content:
                chat_ctx = self.chat_ctx.copy()
                analysis_prompt = (
                    f"استلمت ملف '{stream_info.name}' ({stream_info.mime_type}) من {participant_identity}. "
                    f"هذا ملخص المحتوى:\n\n{file_content}\n\n"
                    f"حلّل المستند وقدّم نقاط قانونية مختصرة ومباشرة."
                )
                chat_ctx.add_message(role="user", content=analysis_prompt)
                await self.update_chat_ctx(chat_ctx)
                await self.session.generate_reply(
                    instructions=(
                        f"تم استلام الملف '{stream_info.name}' ومعالجته. "
                        f"اعطِ ملخصًا واضحًا ثم نقاط قانونية عملية قابلة للتنفيذ."
                    ),
                    allow_interruptions=True,
                )
            else:
                await self.session.generate_reply(
                    instructions=(
                        f"استلمت الملف '{stream_info.name}' لكن نوعه ({stream_info.mime_type}) غير مدعوم. "
                        f"أقدر أعالج PDF والنصوص والصور. جرّب ترفع ملف مدعوم."
                    ),
                    allow_interruptions=True,
                )
        except Exception as e:
            logger.error("❌ [Arabic] error in fallback handling: %s", e, exc_info=True)

