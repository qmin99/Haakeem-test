import logging
import base64
import json
from livekit.agents import Agent
from livekit.plugins import deepgram, groq
from livekit.plugins.azure import TTS as AzureTTS
import re


class ClickToTalkAgent(Agent):
    """Click-to-talk agent that waits for user to finish speaking completely before responding"""

    def __init__(self) -> None:
        super().__init__(
            instructions=(
                "Your name is Haakeem, an AI legal assistant. Your Company is Binfin8. "
                "Users can speak for as long as they want without interruption. "
                "You should use short and concise responses, avoiding usage of unpronounceable punctuation. "
                "Provide thoughtful, comprehensive responses since users may share longer, detailed questions. "
                "Focus on legal guidance, document review, case analysis, or whatever legal assistance they need. "
                "Be thorough in your responses since this agent allows for more in-depth discussion. "
                "Wait for the user to finish speaking completely before responding."
            ),
            stt=deepgram.STT(
                model="nova-2",
                language="en",
                keywords=[
                    ("HAAKEEM", 9.0),
                    ("Haakeem", 6.0),
                    ("Binfin8", 9.0),
                ],

            ),
            llm=groq.LLM(model="llama-3.1-8b-instant"),
            tts=AzureTTS(
                voice="en-US-OnyxTurboMultilingualNeural",
                language="en-US",
            ),
        )

    async def on_enter(self):
        await self.session.generate_reply(
            instructions="Hello! I'm Haakeem, your AI legal assistant. How can I assist you today?",
            allow_interruptions=True,
        )

    async def on_final_transcription(self, text: str) -> str:
        # Normalize brand name variants in English transcripts only
        return self._normalize_brand(text)

    async def _file_received(self, reader, participant_identity):
        logger = logging.getLogger("multi-agent-ptt")
        stream_info = reader.info
        logger.info("📄 [ClickToTalk] received byte stream: %s (%s)", stream_info.name, stream_info.mime_type)
        file_bytes = bytearray()
        async for chunk in reader:
            file_bytes.extend(chunk)
        await self._file_received_fallback(bytes(file_bytes), stream_info, participant_identity)

    async def _process_file(self, file_bytes, stream_info):
        # Delegate to AttorneyAgent-style logic via simple in-file copy for consistency
        logger = logging.getLogger("multi-agent-ptt")
        mime_type = stream_info.mime_type
        file_name = stream_info.name

        try:
            if mime_type == "text/plain":
                try:
                    content = file_bytes.decode("utf-8")
                    return content
                except UnicodeDecodeError:
                    for enc in ["latin-1", "cp1252", "iso-8859-1"]:
                        try:
                            content = file_bytes.decode(enc)
                            return content
                        except UnicodeDecodeError:
                            continue
                    return None
            elif mime_type == "application/pdf":
                try:
                    import io
                    import PyPDF2
                    pdf_stream = io.BytesIO(file_bytes)
                    pdf_reader = PyPDF2.PdfReader(pdf_stream)
                    extracted_text = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    if extracted_text.strip():
                        full_content = f"PDF Document: {file_name}\nContent:\n{extracted_text}"
                        return full_content
                    return (
                        f"PDF document '{file_name}' received but appears to contain no extractable text (may be image-based or encrypted)."
                    )
                except Exception as e:
                    return f"PDF document '{file_name}' received but encountered error during processing: {str(e)}"
            elif mime_type.startswith("image/"):
                image_b64 = base64.b64encode(file_bytes).decode("utf-8")
                return (
                    f"Image file '{file_name}' received ({mime_type}, {len(file_bytes)} bytes). "
                    f"Base64 data available for vision analysis: {image_b64[:100]}...[truncated]"
                )
            elif mime_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ]:
                try:
                    import io
                    from docx import Document
                    doc_stream = io.BytesIO(file_bytes)
                    doc = Document(doc_stream)
                    extracted_text = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            extracted_text += paragraph.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    extracted_text += cell.text + " "
                            extracted_text += "\n"
                    if extracted_text.strip():
                        full_content = f"Word Document: {file_name}\nContent:\n{extracted_text}"
                        return full_content
                    return f"Word document '{file_name}' received but appears to contain no text content."
                except Exception as e:
                    return f"Word document '{file_name}' received but encountered error during processing: {str(e)}"
            elif mime_type == "application/json":
                try:
                    content = json.loads(file_bytes.decode("utf-8"))
                    json_content = f"JSON file '{file_name}' content:\n{json.dumps(content, indent=2)}"
                    return json_content
                except Exception:
                    return None
            else:
                return (
                    f"Received file '{file_name}' with unsupported type '{mime_type}'. Supported types include: "
                    f"text/plain, application/pdf, images, and Word documents."
                )
        except Exception:
            return None

    def _normalize_brand(self, text: str) -> str:
        try:
            pattern = re.compile(r"(?i)(h\s*a\s*a\s*k\s*(?:i|e)?\s*e\s*e\s*m|ha+\s*k[iy]e?m|hakim|hakeem|haakeem|hakem|akim)")
            return pattern.sub("HAAKEEM", text)
        except Exception:
            return text

    async def _file_received_fallback(self, file_bytes, stream_info, participant_identity):
        logger = logging.getLogger("multi-agent-ptt")
        logger.info("📄 [ClickToTalk] fallback upload: %s (%s)", stream_info.name, stream_info.mime_type)
        try:
            file_content = await self._process_file(file_bytes, stream_info)
            if file_content:
                chat_ctx = self.chat_ctx.copy()
                analysis_prompt = (
                    f"I've received a file '{stream_info.name}' ({stream_info.mime_type}) "
                    f"from {participant_identity}. Here's the content:\n\n{file_content}\n\n"
                    f"Please analyze this document and provide legal insights or answer any questions about it."
                )
                chat_ctx.add_message(role="user", content=analysis_prompt)
                await self.update_chat_ctx(chat_ctx)
                await self.session.generate_reply(
                    instructions=(
                        f"You have received and analyzed the file '{stream_info.name}' via upload. "
                        f"Provide a helpful summary and legal analysis of the document content. "
                        f"Be thorough and professional in your response."
                    ),
                    allow_interruptions=True,
                )
            else:
                await self.session.generate_reply(
                    instructions=(
                        f"I received the file '{stream_info.name}', but I wasn't able to process "
                        f"this file type ({stream_info.mime_type}). I can help analyze PDF documents, "
                        f"text files, and images. Please try uploading a supported file format."
                    ),
                    allow_interruptions=True,
                )
        except Exception as e:
            logger.error("❌ [ClickToTalk] error in fallback handling: %s", e, exc_info=True)

