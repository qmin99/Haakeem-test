import logging
import base64
import json
from livekit.agents import Agent
from livekit.plugins import deepgram, groq
from livekit.plugins.azure import TTS as AzureTTS
import re


class AttorneyAgent(Agent):
    """Attorney agent for continuous conversation and legal guidance"""

    def __init__(self) -> None:
        super().__init__(
            instructions=(
                "Your name is Haakeem, an AI legal assistant. Your Company is Binfin8. "
                "You should use short and concise responses, avoiding usage of unpronounceable punctuation. "
                "Be concise and clear in your responses. Focus on legal guidance, document review, case analysis, "
                "or whatever legal assistance they need. When users ask legal questions, provide clear, actionable advice. "
            ),
            stt=deepgram.STT(
                model="nova-2",
                language="en",
                # Deepgram expects list of (keyword, intensifier)
                keywords=[
                    ("HAAKEEM", 9.0),
                    ("Haakeem", 6.0),
                    ("Binfin8", 9.0),
                ],
            ),
            llm=groq.LLM(model="llama-3.1-8b-instant"),
            tts=AzureTTS(
                voice="en-US-DavisNeural",
                language="en-US",
            ),
        )

    async def on_enter(self):
        await self.session.generate_reply(
            instructions="Hello! I'm Haakeem, your AI legal assistant. How can I assist you today?",
            allow_interruptions=True,
        )

    async def on_final_transcription(self, text: str) -> str:
        # Normalize brand name variants in English transcripts only
        return self._normalize_brand(text)

    async def _file_received(self, reader, participant_identity):
        logger = logging.getLogger("multi-agent-ptt")
        logger.info("📄 Received byte stream from %s", participant_identity)
        logger.info("📄 Agent type: %s", self.__class__.__name__)

        try:
            stream_info = reader.info
            logger.info(
                "📄 Stream details - Topic: %s, Name: %s, MIME: %s",
                stream_info.topic,
                stream_info.name,
                stream_info.mime_type,
            )

            file_bytes = bytearray()
            async for chunk in reader:
                file_bytes.extend(chunk)

            file_data = bytes(file_bytes)
            logger.info("📄 Complete file received: %d bytes", len(file_data))

            file_content = await self._process_file(file_data, stream_info)

            if file_content:
                chat_ctx = self.chat_ctx.copy()
                analysis_prompt = (
                    f"I've received a file '{stream_info.name}' ({stream_info.mime_type}) "
                    f"from {participant_identity}. Here's the content:\n\n{file_content}\n\n"
                    f"Please analyze this document and provide legal insights or answer any questions about it."
                )
                chat_ctx.add_message(role="user", content=analysis_prompt)
                await self.update_chat_ctx(chat_ctx)

                logger.info("📄 File processed and added to chat context")

                await self.session.generate_reply(
                    instructions=(
                        f"You have received and analyzed the file '{stream_info.name}'. "
                        f"Provide a helpful summary and legal analysis of the document content. "
                        f"Be thorough and professional in your response."
                    ),
                    allow_interruptions=True,
                )
            else:
                await self.session.generate_reply(
                    instructions=(
                        f"I received the file '{stream_info.name}', but I wasn't able to process "
                        f"this file type ({stream_info.mime_type}). I can help analyze PDF documents, "
                        f"text files, and images. Please try uploading a supported file format."
                    ),
                    allow_interruptions=True,
                )
        except Exception as e:
            logger.error("❌ Error processing byte stream: %s", e, exc_info=True)
            await self.session.generate_reply(
                instructions=(
                    "I encountered an error while processing your file. Please try uploading "
                    "the file again or contact support if the issue persists."
                ),
                allow_interruptions=True,
            )

    async def _process_file(self, file_bytes, stream_info):
        logger = logging.getLogger("multi-agent-ptt")
        mime_type = stream_info.mime_type
        file_name = stream_info.name

        try:
            logger.info("📄 Processing file: %s (%s)", file_name, mime_type)

            if mime_type == "text/plain":
                try:
                    content = file_bytes.decode("utf-8")
                    logger.info("📄 Successfully decoded text file")
                    return content
                except UnicodeDecodeError:
                    for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                        try:
                            content = file_bytes.decode(encoding)
                            logger.info(f"📄 Decoded text file using {encoding}")
                            return content
                        except UnicodeDecodeError:
                            continue
                    logger.warning("📄 Could not decode text file with any encoding")
                    return None

            elif mime_type == "application/pdf":
                logger.info("📄 PDF file detected - extracting content")
                try:
                    import io
                    import PyPDF2

                    pdf_stream = io.BytesIO(file_bytes)
                    pdf_reader = PyPDF2.PdfReader(pdf_stream)

                    extracted_text = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

                    if extracted_text.strip():
                        logger.info(
                            f"📄 Successfully extracted {len(extracted_text)} characters from PDF"
                        )
                        full_content = f"PDF Document: {file_name}\nContent:\n{extracted_text}"
                        return full_content
                    else:
                        logger.warning(
                            "📄 PDF appears to be empty or contains only images"
                        )
                        return (
                            f"PDF document '{file_name}' received but appears to contain no extractable "
                            f"text (may be image-based or encrypted)."
                        )
                except ImportError:
                    logger.error("📄 PyPDF2 not installed - cannot process PDF files")
                    return (
                        f"PDF document '{file_name}' received ({len(file_bytes)} bytes) but PyPDF2 "
                        f"library is not available for text extraction."
                    )
                except Exception as e:
                    logger.error(f"📄 Error processing PDF: {e}")
                    return (
                        f"PDF document '{file_name}' received but encountered error during processing: {str(e)}"
                    )

            elif mime_type.startswith("image/"):
                logger.info("📄 Image file detected")
                image_b64 = base64.b64encode(file_bytes).decode("utf-8")
                return (
                    f"Image file '{file_name}' received ({mime_type}, {len(file_bytes)} bytes). "
                    f"Base64 data available for vision analysis: {image_b64[:100]}...[truncated]"
                )

            elif mime_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ]:
                logger.info("📄 Word document detected - extracting content")
                try:
                    import io
                    from docx import Document

                    doc_stream = io.BytesIO(file_bytes)
                    doc = Document(doc_stream)

                    extracted_text = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            extracted_text += paragraph.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    extracted_text += cell.text + " "
                            extracted_text += "\n"

                    if extracted_text.strip():
                        logger.info(
                            f"📄 Successfully extracted {len(extracted_text)} characters from Word document"
                        )
                        full_content = f"Word Document: {file_name}\nContent:\n{extracted_text}"
                        return full_content
                    else:
                        logger.warning("📄 Word document appears to be empty")
                        return (
                            f"Word document '{file_name}' received but appears to contain no text content."
                        )
                except ImportError:
                    logger.error(
                        "📄 python-docx not installed - cannot process Word documents"
                    )
                    return (
                        f"Word document '{file_name}' received ({len(file_bytes)} bytes) but python-docx "
                        f"library is not available for text extraction."
                    )
                except Exception as e:
                    logger.error(f"📄 Error processing Word document: {e}")
                    return (
                        f"Word document '{file_name}' received but encountered error during processing: {str(e)}"
                    )

            elif mime_type == "application/json":
                try:
                    content = json.loads(file_bytes.decode("utf-8"))
                    logger.info("📄 Successfully parsed JSON file")
                    json_content = f"JSON file '{file_name}' content:\n{json.dumps(content, indent=2)}"
                    return json_content
                except (json.JSONDecodeError, UnicodeDecodeError) as e:
                    logger.warning(f"📄 Failed to parse JSON: {e}")
                    return None

            else:
                logger.warning("📄 Unsupported file type: %s", mime_type)
                return (
                    f"Received file '{file_name}' with unsupported type '{mime_type}'. Supported types include: "
                    f"text/plain, application/pdf, images, and Word documents."
                )
        except Exception as e:
            logger.error("❌ Error processing file content: %s", e, exc_info=True)
            return None

    def _normalize_brand(self, text: str) -> str:
        try:
            pattern = re.compile(r"(?i)(h\s*a\s*a\s*k\s*(?:i|e)?\s*e\s*e\s*m|ha+\s*k[iy]e?m|hakim|hakeem|haakeem|hakem|akim)")
            return pattern.sub("HAAKEEM", text)
        except Exception:
            return text

    async def _file_received_fallback(self, file_bytes, stream_info, participant_identity):
        logger = logging.getLogger("multi-agent-ptt")
        logger.info(
            "📄 Processing file upload fallback from %s: '%s'",
            participant_identity,
            stream_info.name,
        )
        try:
            file_content = await self._process_file(file_bytes, stream_info)

            if file_content:
                chat_ctx = self.chat_ctx.copy()
                analysis_prompt = (
                    f"I've received a file '{stream_info.name}' ({stream_info.mime_type}) "
                    f"from {participant_identity}. Here's the content:\n\n{file_content}\n\n"
                    f"Please analyze this document and provide legal insights or answer any questions about it."
                )
                chat_ctx.add_message(role="user", content=analysis_prompt)
                await self.update_chat_ctx(chat_ctx)

                logger.info("📄 File processed and added to chat context (fallback)")

                await self.session.generate_reply(
                    instructions=(
                        f"You have received and analyzed the file '{stream_info.name}' via upload. "
                        f"Provide a helpful summary and legal analysis of the document content. "
                        f"Be thorough and professional in your response."
                    ),
                    allow_interruptions=True,
                )
            else:
                await self.session.generate_reply(
                    instructions=(
                        f"I received the file '{stream_info.name}', but I wasn't able to process "
                        f"this file type ({stream_info.mime_type}). I can help analyze PDF documents, "
                        f"text files, and images. Please try uploading a supported file format."
                    ),
                    allow_interruptions=True,
                )
        except Exception as e:
            logger.error("❌ Error processing file upload fallback: %s", e, exc_info=True)
            await self.session.generate_reply(
                instructions=(
                    "I encountered an error while processing your uploaded file. Please try uploading "
                    "the file again or contact support if the issue persists."
                ),
                allow_interruptions=True,
            )


